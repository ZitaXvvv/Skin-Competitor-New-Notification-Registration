"""Convert docs/User_Guide.md and docs/Admin_Guide.md to .docx"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR = Path(__file__).parent.parent / "docs"

def _inline_runs(para, text):
    """Add runs to para, handling **bold** and `code` inline markup."""
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for p in parts:
        if p.startswith('**') and p.endswith('**'):
            run = para.add_run(p[2:-2])
            run.bold = True
        elif p.startswith('`') and p.endswith('`'):
            run = para.add_run(p[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            run = para.add_run(p)

def _shade_para(para, fill='F5F5F5'):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def _shade_cell(cell, fill='E8E8E8'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def md_to_docx(md_path: Path, out_path: Path):
    doc = Document()

    # Default font
    for st_name in ('Normal', 'Default Paragraph Font'):
        try:
            st = doc.styles[st_name]
            st.font.name = 'Arial'
        except Exception:
            pass

    lines = md_path.read_text(encoding='utf-8').splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            return
        ncols = max(len(r) for r in table_rows)
        tbl = doc.add_table(rows=len(table_rows), cols=ncols)
        tbl.style = 'Table Grid'
        for ri, row in enumerate(table_rows):
            for ci in range(ncols):
                cell_text = row[ci] if ci < len(row) else ''
                cell = tbl.cell(ri, ci)
                cell.text = ''
                para = cell.paragraphs[0]
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                _inline_runs(para, cell_text)
                for run in para.runs:
                    run.font.size = Pt(9)
                    if ri == 0:
                        run.bold = True
                if ri == 0:
                    _shade_cell(cell, 'D9E1F2')
        table_rows = []
        in_table = False

    while i < len(lines):
        line = lines[i]

        # ── code block ──
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                para = doc.add_paragraph()
                run = para.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(8.5)
                para.paragraph_format.left_indent = Inches(0.4)
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after = Pt(4)
                _shade_para(para)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── markdown table ──
        if re.match(r'^\s*\|', line):
            if not in_table:
                in_table = True
                table_rows = []
            # skip separator row
            if re.match(r'^\s*\|[\s\-|:]+\|\s*$', line):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # ── blank line ──
        if not line.strip():
            i += 1
            continue

        # ── headings ──
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=min(level, 4))
            i += 1
            continue

        # ── blockquote ──
        if line.startswith('> '):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.35)
            _inline_runs(para, line[2:])
            for run in para.runs:
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                run.italic = True
            i += 1
            continue

        # ── checkbox list ──
        if re.match(r'^- \[ \] ', line):
            para = doc.add_paragraph(style='List Bullet')
            _inline_runs(para, '☐  ' + line[6:].strip())
            i += 1
            continue

        # ── bullet list ──
        if re.match(r'^[-*] ', line):
            para = doc.add_paragraph(style='List Bullet')
            _inline_runs(para, line[2:].strip())
            i += 1
            continue

        # ── indented bullet ──
        if re.match(r'^  [-*] ', line):
            para = doc.add_paragraph(style='List Bullet 2')
            _inline_runs(para, line[4:].strip())
            i += 1
            continue

        # ── numbered list ──
        if re.match(r'^\d+\. ', line):
            para = doc.add_paragraph(style='List Number')
            _inline_runs(para, re.sub(r'^\d+\. ', '', line))
            i += 1
            continue

        # ── horizontal rule ──
        if re.match(r'^[-*_]{3,}\s*$', line):
            doc.add_paragraph('─' * 55)
            i += 1
            continue

        # ── normal paragraph ──
        para = doc.add_paragraph()
        _inline_runs(para, line.strip())
        i += 1

    # flush any trailing table
    if in_table:
        flush_table()

    doc.save(str(out_path))
    print(f'✓  {out_path.name}')


if __name__ == '__main__':
    md_to_docx(DOCS_DIR / 'User_Guide.md',  DOCS_DIR / 'User_Guide.docx')
    md_to_docx(DOCS_DIR / 'Admin_Guide.md', DOCS_DIR / 'Admin_Guide.docx')
    print('Done.')
